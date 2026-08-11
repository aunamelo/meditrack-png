"""Generate Project Report 1 as a Word document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)


def add_heading_centered(text, level=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 0 else 12)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_para(text, bold=False, center=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


# ---------- Cover ----------
add_heading_centered("Department of Information Systems")
add_heading_centered("Faculty of Business & Informatics")
add_heading_centered("Divine Word University")
doc.add_paragraph()
add_heading_centered("Project Report 1")
doc.add_paragraph()
add_para(
    "MediTrack PNG — A Web-Based Medicine Supply Chain Tracking System for Papua New Guinea",
    bold=True,
    center=True,
)
doc.add_paragraph()
add_para("Keith Banks Pala", center=True)
add_para("Bachelor of Information Systems", center=True)
add_para("2026", center=True)

doc.add_page_break()

# ---------- TOC ----------
add_para("Table of Contents", bold=True)
add_para("1. Introduction\t1")
add_para("2. Work achievements\t1")
add_para("3. Issues faced\t1")
add_para("4. Next plan of action\t1")
add_para("5. Status\t2")

doc.add_page_break()

# ---------- 1 ----------
doc.add_heading("1. Introduction", level=1)
add_para(
    "This report presents the progress of the MediTrack PNG final-year project for Weeks 1–6 "
    "of the approved project schedule. It summarises work completed, challenges encountered, "
    "the current project status, and planned activities for Weeks 7–12."
)
add_para(
    "MediTrack PNG is a web-based drug tracking system designed to improve visibility and "
    "control across Papua New Guinea’s public medicine supply chain. Phase 1 is piloted for "
    "Madang Province, following medicines from the National Department of Health (NDoH) through "
    "the Lae Area Medical Store (Lae AMS) to Modilon General Hospital. The system addresses gaps "
    "in manual, disconnected processes that contribute to stockouts, expiry risk, and weak oversight."
)
add_para(
    "The project is built primarily with Laravel (PHP), MySQL, Blade, Tailwind CSS, and Spatie "
    "Laravel Permission for role-based access across five portals: NDoH Admin, Procurement Officer, "
    "Store Manager, Pharmacy Manager, and Pharmacist."
)
add_para(
    "This report covers progress and deliverables for Weeks 1–6, issues and how they were managed, "
    "the next plan of action for Weeks 7–12, and an evidence-based statement of overall status."
)

# ---------- 2 ----------
doc.add_heading("2. Work achievements", level=1)

doc.add_heading("2.1 Weeks 1–3 — Requirements, analysis, and design", level=2)

table = doc.add_table(rows=4, cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Week", "Phase (schedule)", "Activities completed", "Deliverables"]
for i, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True)

rows = [
    [
        "1",
        "Requirements",
        "Stakeholder/problem framing around NDoH → Lae AMS → Modilon; scope definition (Phase 1 Madang pilot); finalisation of project proposal",
        "Project proposal; problem statement; in-scope / out-of-scope list; high-level objectives",
    ],
    [
        "2",
        "Analysis",
        "Identification of actors and use cases for five roles; mapping of procurement, transfer, and hospital workflows",
        "Use case diagrams; activity diagrams; system sequence diagrams (SSD) for key flows (login, procurement order, stock transfer)",
    ],
    [
        "3",
        "Design",
        "Conceptual and logical data model; supply-chain levels (ndoh, lae_ams, modilon_hospital); architecture of the web application and deployment approach",
        "ERD / database schema design; system architecture notes; supporting design documentation",
    ],
]
for r_i, row in enumerate(rows, start=1):
    for c_i, val in enumerate(row):
        set_cell_text(table.rows[r_i].cells[c_i], val)

add_para(
    "Evidence: proposal and Week 1–3 diagrams in the assigned shared drive; Trello cards for "
    "requirements, use cases, and ERD marked Done. (Insert Drive and Trello links before submission.)"
)

doc.add_heading("2.2 Weeks 4–5 — Development 1: environment, authentication, and roles", level=2)
add_para(
    "Aligned with the schedule (Set up development environment, implement user authentication and role management)."
)
add_para("Major tasks completed:", bold=True)
add_bullets(
    [
        "Set up the local development environment (PHP, Composer, Laravel, MySQL, Vite/Node).",
        "Initialised Git and pushed work to the assigned GitHub repository.",
        "Implemented authentication (login / session).",
        "Implemented role-based access control for five roles using Spatie permissions.",
        "Built role-scoped dashboards and navigation so each user only sees modules for their level.",
        "Implemented the guest/home “Choose your portal role” entry into role-specific login paths.",
    ]
)
add_para("System components delivered:", bold=True)
add_bullets(
    [
        "User accounts and roles: Admin, Procurement Officer, Store Manager, Pharmacy Manager, Pharmacist.",
        "Portal configuration as a single source of truth for role labels and inventory levels.",
        "Protected routes and middleware-based authorisation.",
    ]
)
add_para(
    "Evidence: GitHub commits for auth/RBAC/dashboards; screenshots of login, role selection, and "
    "role dashboards; Trello cards for environment, auth, and RBAC marked Done. Repository: "
    "https://github.com/aunamelo/meditrack-png (add lecturer remote link if required)."
)

doc.add_heading("2.3 Week 6 — Development 2: inventory and procurement (ahead of baseline)", level=2)
add_para(
    "The schedule places inventory and procurement across Weeks 6–7. By the end of Week 6, this "
    "work was substantially advanced, which is why the project is ahead of schedule."
)
add_para("Major tasks completed:", bold=True)
add_bullets(
    [
        "Implemented the NDoH medicine catalog (procurement reference list separate from physical stock).",
        "Implemented inventory batches with supply-chain level, batch number, quantity on hand, and expiry.",
        "Progressed procurement orders (orders / line items) for sourcing and receiving stock into NDoH.",
        "Continued schema migrations supporting catalog, suppliers, and inventory linkages.",
        "Separated catalog medicines from warehouse batches in the data model to avoid redesign later.",
    ]
)
add_para("Deliverables by end of Week 6:", bold=True)
add_bullets(
    [
        "Working multi-role authenticated application.",
        "Usable foundation of inventory and procurement modules (beyond a “Week 6 start only” baseline).",
        "Updated ERD documentation reflecting catalog vs batch separation.",
    ]
)
add_para(
    "Evidence: migrations/models/controllers for medicines, drugs, and orders on GitHub; screenshots "
    "of catalog and inventory screens; Trello Development 2 cards moved to Done / near-complete as applicable."
)

doc.add_heading("2.4 Summary of Weeks 1–6 deliverables", level=2)
add_bullets(
    [
        "Approved project proposal and Madang Phase 1 scope.",
        "Analysis diagrams (use case, activity, SSD).",
        "Design artefacts (ERD, architecture).",
        "Running Laravel application with authentication and five role portals.",
        "Strong progress on inventory and procurement (Development 2), completed earlier than the Week 7 finish target.",
        "Trello, shared drive, and GitHub kept consistent with reported progress.",
    ]
)
add_para(
    "Note: Insert screenshots and live links (Trello, Drive, GitHub commits) before submission."
)

# ---------- 3 ----------
doc.add_heading("3. Issues faced", level=1)

doc.add_heading("3.1 Environment and tooling setup (Weeks 4–5)", level=2)
add_para(
    "What happened: Initial Laravel/Vite/MySQL setup on Windows required several configuration "
    "steps (PHP extensions, Node/Vite, .env database credentials)."
)
add_para("Impact: Delayed the first successful local run of the application.")
add_para(
    "Action taken: Standardised .env, verified php artisan serve and frontend build tooling, "
    "and documented setup for reuse."
)
add_para("Status: Resolved.")

doc.add_heading("3.2 Role separation and route design (Weeks 4–5)", level=2)
add_para(
    "What happened: Designing five portals without leaking navigation or data across roles was "
    "more complex than a single shared dashboard."
)
add_para("Impact: Risk of incorrect access and confusing UX if routes were not role-scoped.")
add_para(
    "Action taken: Centralised role metadata; used Spatie roles and role-specific dashboards; "
    "tested login for each role."
)
add_para("Status: Resolved for core auth/RBAC; continues to be checked as modules are added.")

doc.add_heading("3.3 Modelling catalog medicines vs physical stock batches (Week 6)", level=2)
add_para(
    "What happened: Early risk of treating “medicine type” and “warehouse batch” as one concept, "
    "which would break multi-level inventory and procurement."
)
add_para("Impact: Could have forced a costly redesign of procurement and transfers later.")
add_para(
    "Action taken: Separated medicines (catalog) from drugs (batches at NDoH / Lae AMS / Modilon) "
    "and documented this in the ERD."
)
add_para("Status: Resolved.")
add_para(
    "These issues were managed without pushing the project behind the Week 6 checkpoint; remaining "
    "Development 2 polish fits within Week 7."
)

# ---------- 4 ----------
doc.add_heading("4. Next plan of action (Weeks 7–12)", level=1)

doc.add_heading("4.1 Planned work", level=2)
table2 = doc.add_table(rows=6, cols=4)
table2.style = "Table Grid"
headers2 = ["Weeks", "Phase", "Planned tasks & deliverables", "Target"]
for i, h in enumerate(headers2):
    set_cell_text(table2.rows[0].cells[i], h, bold=True)

plan_rows = [
    [
        "7",
        "Development 2 (finish / buffer)",
        "Finalise remaining procurement UI and stock-receive flows; harden inventory views",
        "End of Week 7",
    ],
    [
        "8–9",
        "Development 3",
        "NDoH→Lae AMS shipments; Lae AMS→Modilon road deliveries; dispensing records at Modilon",
        "End of Week 9",
    ],
    [
        "10",
        "Development 4",
        "Rule-based alerts (low stock, near expiry); role reporting dashboards",
        "End of Week 10",
    ],
    [
        "11",
        "Testing",
        "Unit/feature tests; full-corridor integration; UAT with role scenarios",
        "End of Week 11",
    ],
    [
        "12",
        "Deployment & docs",
        "Deploy to assigned host; documentation; Information Systems Symposium presentation",
        "Symposium date",
    ],
]
for r_i, row in enumerate(plan_rows, start=1):
    for c_i, val in enumerate(row):
        set_cell_text(table2.rows[r_i].cells[c_i], val)

doc.add_heading("4.2 Features to develop next", level=2)
add_bullets(
    [
        "Complete any remaining procurement/inventory polish (Week 7).",
        "Distribution tracking (national freight + hospital road leg).",
        "Dispensing and patient-linked records.",
        "Alerts and reporting dashboards.",
        "Testing, documentation, and deployment.",
    ]
)

doc.add_heading("4.3 Updated WBS (Weeks 7–12)", level=2)
add_bullets(
    [
        "W7 — Close Development 2",
        "W8–9 — Distribution & dispensing",
        "W10 — Alerts & reporting",
        "W11 — Testing & UAT",
        "W12 — Deploy + symposium pack",
    ]
)
add_para(
    "Trello will be updated with due dates, priorities, and milestones for these weeks. Shared "
    "drive will hold UAT scripts and presentation drafts; GitHub will continue to receive regular commits."
)

# ---------- 5 ----------
doc.add_heading("5. Status", level=1)
add_para("Current overall status: Ahead of schedule", bold=True)
add_para("Why:")
add_bullets(
    [
        "Weeks 1–3 deliverables (proposal, analysis diagrams, ERD/architecture) are complete.",
        "Weeks 4–5 authentication and five-portal role management are implemented and usable.",
        "Week 6 was scheduled as the start of Development 2 (inventory and procurement). By the end of Week 6, inventory and procurement work was already well advanced, putting completion of Development 2 earlier than the Week 7 target on the approved schedule.",
        "A running local application and GitHub history provide concrete evidence of this progress, alongside updated Trello cards and shared-drive artefacts.",
    ]
)
add_para(
    "Being ahead of schedule provides buffer time in Week 7 to polish procurement/inventory, then "
    "move into Development 3 (distribution and dispensing) without compressing later testing and "
    "deployment weeks."
)

out = r"c:\Users\ID230178\Desktop\meditrack-png\docs\Project_Report_1_Keith_Banks_Pala.docx"
doc.save(out)
print(out)
